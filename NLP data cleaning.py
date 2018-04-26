import sys
from docx import Document
import nltk
import re
import json
# sys.path.append('.\\nlpAbbreviation.py')

#method to find entity group but not accurate
def extract_entities(clean_sentence):
	for i in range(len(clean_sentence)):
		print('---------------------------------------')
		text = str(clean_sentence[i])
		print(text)
		for sent in nltk.sent_tokenize(text):
			for chunk in nltk.ne_chunk(nltk.pos_tag(nltk.word_tokenize(sent))):
				# print(nltk.word_tokenize(sent))
				# print(nltk.pos_tag(nltk.word_tokenize(sent)))
				# print(nltk.ne_chunk(nltk.pos_tag(nltk.word_tokenize(sent))))#chunk return nested nltk.tree.Tree object so you have to traverse the Tree object to get NEs
				if hasattr(chunk, 'label'):
					print(chunk.label(), ' '.join(c[0] for c in chunk.leaves()))

def remove_like_reply(sentence):
	clean_sentence = []
	pattern1 = '^Repl'
	for i in range(len(sentence)):
		sentence[i] = sentence[i].split(' ')
		if(len(sentence[i]) <= 3 or (sentence[i][0] == 'Like' and sentence[i][2] == 'Reply')):
			#ignore sentence left than 4 words
			pass
		elif((re.match(pattern1,sentence[i][1]) != None)==True):
			pass
		else:
			clean_sentence.append(sentence[i])
		
	return clean_sentence

def remove_emoji(clean_sentence):
	emoji_pattern = re.compile("["
		u"\U00000021-\U0000002F" #symbols !"#$%^.
		u"\U0001F600-\U0001F64F"  # emoticons
        u"\U0001F300-\U0001F5FF"  # symbols & pictographs
        u"\U0001F680-\U0001F6FF"  # transport & map symbols
        u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
                           "]+", flags=re.UNICODE)
	for i in range(len(clean_sentence)):
		# print()
		# text = str(clean_sentence[i])
		for j in range(len(clean_sentence[i])):
			clean_sentence[i][j]= emoji_pattern.sub(r' ',clean_sentence[i][j])	
			print('xxxxxx '+clean_sentence[i][j])
	return clean_sentence

def remove_wordEmoticon(clean_sentence):
	for sentence in clean_sentence:
		word_index_to_be_delete = []
		for i in range(len(sentence)):
			sentence[i] = " ".join(sentence[i].split())
			if(sentence[i].lower() == "emoticon"):
				sentence[i]=""
				sentence[i-1]=""
	return clean_sentence



def removeNumtoWords(clean_sentence):
	pattern = re.compile(r'([A-z]){1,}\d+')
	for sentence in clean_sentence:
		for i in range(len(sentence)):
			try:
				if(pattern.match(sentence[i])):
					new_word = ''
					for j in range(int(sentence[i][-1])):
						new_word = new_word +' '+sentence[i][:-1]

					sentence[i] = new_word
			except:
				print(sentence[i]+' gt error.')

	return clean_sentence

# f = lambda x: x*2
# f(2)
#map ~ list.append
#parallel processing, faster than forloop list.append

def joinTokens2Sentence(tokens):
	deto_clean_sentence = list(map(lambda token: ' '.join(token), tokens))
	# deto_sentence = []
	# for token in tokens:
	# 	deto_sentence.append(' '.join(token)) 
	return deto_clean_sentence

def splitSentence2Tokens(sentence):
	for i in range(len(sentence)):
		tokens.append(sentence[i].split(' '))
	return tokens

document = Document('.\\NLP part1.docx')
sentence = []
tokens = []
clean_sentence = []
for para in document.paragraphs:
	#para.text.replaceAll("\n","\s")
	sentence.append(para.text)
	# print(tokens)
print(sentence)
clean_sentence = remove_like_reply(sentence)
# extract_entities(clean_sentence)
clean_sentence = remove_emoji(clean_sentence)
clean_sentence = joinTokens2Sentence(clean_sentence)
clean_sentence = splitSentence2Tokens(clean_sentence)
clean_sentence = remove_wordEmoticon(clean_sentence)
clean_sentence = removeNumtoWords(clean_sentence)
clean_sentence = joinTokens2Sentence(clean_sentence)
clean_sentence = splitSentence2Tokens(clean_sentence)
deto_clean_sentence = joinTokens2Sentence(clean_sentence)


new_document = Document()
# write clean sentence to word doc
for i in range(len(deto_clean_sentence)):
	print(deto_clean_sentence[i])
	new_document.add_paragraph(deto_clean_sentence[i])
new_document.save('nlp cleaned part1.docx')

# # for i in range(len(clean_sentence)):
# # 	tokens += nltk.word_tokenize(clean_sentence[i])
# # text = nltk.Text(tokens)
# # print(text.collocations())



